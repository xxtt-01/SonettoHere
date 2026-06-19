"""Tool: doc_reader — Word 文档阅读。"""

import os

from pydantic import BaseModel, Field

from tools.base import (
    ToolBase,
    check_path_whitelisted,
    check_sonetto_blocker,
    format_error,
    format_success,
)


class DocReaderInput(BaseModel):
    get_doc: bool = Field(
        default=False, description="设为 true 以获取使用说明和领域知识"
    )
    operation: str = Field(
        default="",
        description="操作: get_metadata/extract_text/search_text/get_paragraphs/get_tables",
    )
    file_path: str = Field(default="", description="DOCX 文件路径")
    start_paragraph: int = Field(default=0, description="起始段落索引（从0开始）")
    end_paragraph: int | None = Field(
        default=None, description="结束段落索引（从0开始）"
    )
    query: str = Field(default="", description="搜索关键词（search_text 操作）")
    case_sensitive: bool = Field(default=False, description="搜索是否区分大小写")
    max_length: int = Field(default=10000, description="文本最大长度")


class DocReaderTool(ToolBase):
    name: str = "doc_reader"
    description: str = (
        "读取 Word 文档（.docx）：元数据、文本提取、段落、表格、关键词搜索。"
        "仅支持 .docx，不支持旧 .doc 格式。[调用积极性: 可自由看情况调用] [get_doc: 仅在发生错误时 get_doc]"
    )
    args_schema: type[BaseModel] = DocReaderInput

    def _run(
        self,
        get_doc: bool = False,
        operation: str = "",
        file_path: str = "",
        start_paragraph: int = 0,
        end_paragraph: int | None = None,
        query: str = "",
        case_sensitive: bool = False,
        max_length: int = 10000,
    ) -> str:
        if get_doc:
            return self._load_doc()
        if not operation:
            return format_error("operation 不能为空")
        if not file_path:
            return format_error("file_path 不能为空")

        # ── SonettoBlocker 安全检查 ────────────────────────────────
        blocked = check_sonetto_blocker(file_path)
        if blocked:
            return format_error(
                "🚫 安全阻断：操作已被 SonettoBlocker 阻断。\n"
                f'在目录 "{blocked}" 中发现了 SonettoBlocker 文件。\n\n'
                "请立即停止当前任务，先说明你为什么需要访问该路径，"
                "再说明下一步打算做什么。"
            )
        # ────────────────────────────────────────────────────────────

        # ── 路径白名单检查 ──────────────────────────────────────────
        blocked = check_path_whitelisted(file_path)
        if blocked:
            return format_error(blocked)
        # ────────────────────────────────────────────────────────────

        if not os.path.exists(file_path):
            return format_error(f"文件不存在: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in (".doc", ".docx"):
            return format_error(f"不是 Word 文档: {file_path}")
        if ext == ".doc":
            return format_error(".doc 格式不支持，请先转换为 .docx")

        try:
            import docx

            if operation == "get_metadata":
                return self._get_metadata(docx, file_path)
            elif operation == "extract_text":
                return self._extract_text(
                    docx, file_path, start_paragraph, end_paragraph, max_length
                )
            elif operation == "search_text":
                return self._search_text(docx, file_path, query, case_sensitive)
            elif operation == "get_paragraphs":
                return self._get_paragraphs(
                    docx, file_path, start_paragraph, end_paragraph
                )
            elif operation == "get_tables":
                return self._get_tables(docx, file_path)
            else:
                return format_error(f"未知操作: {operation}")
        except ImportError:
            return format_error("python-docx 未安装，请运行: pip install python-docx")
        except Exception as e:
            return format_error(f"Word 文档处理失败: {e}")

    def _get_metadata(self, docx, file_path: str) -> str:
        doc = docx.Document(file_path)
        cp = doc.core_properties
        meta = {
            "title": cp.title or "",
            "author": cp.author or "",
            "subject": cp.subject or "",
            "created": str(cp.created) if cp.created else "",
            "modified": str(cp.modified) if cp.modified else "",
            "last_modified_by": cp.last_modified_by or "",
            "keywords": cp.keywords or "",
            "category": cp.category or "",
        }
        return format_success(
            {
                "metadata": meta,
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
                "file_path": os.path.abspath(file_path),
            }
        )

    def _extract_text(
        self, docx, file_path: str, start: int, end: int | None, max_length: int
    ) -> str:
        doc = docx.Document(file_path)
        total = len(doc.paragraphs)

        if end is None or end >= total:
            end = total - 1
        if start >= total:
            return format_error(f"起始段落超出范围: {start}，总段落数: {total}")

        parts = [
            doc.paragraphs[i].text
            for i in range(start, end + 1)
            if doc.paragraphs[i].text.strip()
        ]
        text = "\n\n".join(parts)
        if len(text) > max_length:
            text = text[:max_length] + "\n\n[内容已截断]"

        return format_success(
            {
                "text": text,
                "paragraph_range": [start, end],
                "total_paragraphs": total,
                "text_length": len(text),
            }
        )

    def _search_text(
        self, docx, file_path: str, query: str, case_sensitive: bool
    ) -> str:
        if not query:
            return format_error("请提供搜索关键词")

        doc = docx.Document(file_path)
        q = query if case_sensitive else query.lower()
        results = []

        for idx, para in enumerate(doc.paragraphs):
            if not para.text:
                continue
            check = para.text if case_sensitive else para.text.lower()
            if q in check:
                results.append(
                    {
                        "paragraph_index": idx,
                        "paragraph_number": idx + 1,
                        "content": para.text.strip(),
                    }
                )

        return format_success(
            {
                "query": query,
                "total_paragraphs": len(doc.paragraphs),
                "results": results,
                "total_matches": len(results),
            }
        )

    def _get_paragraphs(self, docx, file_path: str, start: int, end: int | None) -> str:
        doc = docx.Document(file_path)
        total = len(doc.paragraphs)

        if end is None or end >= total:
            end = total - 1
        if start >= total:
            return format_error(f"起始段落超出范围: {start}，总段落数: {total}")

        paras = []
        for i in range(start, end + 1):
            p = doc.paragraphs[i]
            paras.append(
                {
                    "index": i,
                    "number": i + 1,
                    "text": p.text,
                    "style": p.style.name if p.style else None,
                }
            )

        return format_success(
            {
                "paragraphs": paras,
                "paragraph_range": [start, end],
                "total_paragraphs": total,
            }
        )

    def _get_tables(self, docx, file_path: str) -> str:
        doc = docx.Document(file_path)
        tables = []
        for idx, table in enumerate(doc.tables):
            data = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(
                {
                    "index": idx,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": data,
                }
            )

        return format_success({"tables": tables, "total_tables": len(tables)})
